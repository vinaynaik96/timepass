from docx import Document
import pandas as pd

# Create a new Document
doc = Document()
doc.add_heading('Specifications for Running and Training Large Language Models', 0)

# Add content for running a pre-trained model
doc.add_heading('1. Running a Pre-trained Large Language Model (e.g., GPT-4):', level=1)
doc.add_heading('GPU:', level=2)
doc.add_paragraph(
    "A dedicated GPU, preferably from NVIDIA with CUDA compatibility, is essential. "
    "The larger the model, the more VRAM you'll need. For example, a model like GPT-2's "
    "smaller variants can be run on GPUs with 8GB VRAM, but for GPT-3 or GPT-4, GPUs with "
    "16GB or more VRAM (like NVIDIA's A100 or V100) are preferable."
)
doc.add_heading('RAM:', level=2)
doc.add_paragraph("At least 16GB, but 32GB or more is recommended for smoother operations.")
doc.add_heading('Storage:', level=2)
doc.add_paragraph(
    "SSDs are preferred over HDDs for faster data access. The exact storage requirement will "
    "depend on the model size. For instance, GPT-3's 175B parameter model requires hundreds "
    "of gigabytes just for the model weights."
)
doc.add_heading('Software:', level=2)
doc.add_paragraph(
    "Most implementations will require Python with libraries like TensorFlow or PyTorch. "
    "CUDA and cuDNN installations are required for GPU acceleration."
)

# Add content for retraining or fine-tuning a model
doc.add_heading('2. Retraining or Fine-tuning a Large Language Model:', level=1)
doc.add_heading('GPU:', level=2)
doc.add_paragraph(
    "Multiple high-end GPUs or even TPUs are typically required. Training large models can take "
    "weeks or even months on clusters of GPUs. The more VRAM and faster the GPU, the better. "
    "Multi-GPU setups or cloud-based clusters are commonly used."
)
doc.add_heading('RAM:', level=2)
doc.add_paragraph("64GB or more. Distributed setups should have significant RAM on each machine.")
doc.add_heading('Storage:', level=2)
doc.add_paragraph(
    "Several terabytes of SSD storage might be needed, especially if you're working with large datasets."
)
doc.add_heading('Software:', level=2)
doc.add_paragraph(
    "Distributed training setups often require specific software configurations. Tools like Horovod "
    "can be used to manage distributed training with TensorFlow or PyTorch. CUDA and cuDNN are essential."
)
doc.add_heading('Dataset:', level=2)
doc.add_paragraph(
    "The size and quality of the dataset are crucial. For retraining from scratch, multi-terabyte datasets "
    "are used. For fine-tuning, the dataset can be smaller."
)
doc.add_heading('Infrastructure:', level=2)
doc.add_paragraph(
    "Efficient cooling, power backups, and high-speed internet (for cloud-based setups) are essential. "
    "Training these models generates a lot of heat, and interruptions can be costly."
)

# Considerations for retraining vs. fine-tuning
doc.add_heading('3. Considerations for Retraining vs. Fine-tuning:', level=1)
doc.add_paragraph(
    "Retraining from scratch requires a much larger dataset and is computationally more intensive. "
    "The infrastructure needs are considerably higher. Fine-tuning is less resource-intensive as you're "
    "often training on a smaller, task-specific dataset and you're not training the model from scratch. "
    "However, it still requires a powerful GPU setup."
)

# Save the document to a file
doc_file_path = "/mnt/data/large_language_model_specs.docx"
doc.save(doc_file_path)

# Create an Excel file with the same information
data = {
    'Category': ['GPU', 'RAM', 'Storage', 'Software', 'Dataset', 'Infrastructure'],
    'Running a Pre-trained Model': [
        "Dedicated GPU (e.g., NVIDIA with CUDA). VRAM depends on model size.",
        "At least 16GB (32GB recommended)",
        "SSD preferred. Storage depends on model size.",
        "Python, TensorFlow/PyTorch, CUDA, cuDNN",
        "N/A",
        "N/A"
    ],
    'Retraining/Fine-tuning': [
        "Multiple high-end GPUs/TPUs. Multi-GPU setups or cloud clusters.",
        "64GB or more.",
        "Several TBs of SSD storage.",
        "Distributed training tools, TensorFlow/PyTorch, CUDA, cuDNN",
        "Large datasets (multi-TB for retraining). Smaller for fine-tuning.",
        "Cooling, power backups, high-speed internet"
    ]
}

df = pd.DataFrame(data)
excel_file_path = "/mnt/data/large_language_model_specs.xlsx"
df.to_excel(excel_file_path, index=False)

doc_file_path, excel_file_path
